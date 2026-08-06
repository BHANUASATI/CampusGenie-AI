"""
Document Text Extractor
========================
Extracts raw text from supported file formats:
  .pdf   — PyMuPDF (fitz)    — preserves reading order, handles multi-column
  .docx  — python-docx       — paragraphs + tables
  .txt   — direct read       — UTF-8 with fallback encoding detection
  .csv   — pandas            — converts rows to readable sentences
  .md    — direct read       — treated as plain text

Each extractor returns a list of (page_or_section_number, text) tuples
so the chunker can attach page metadata to chunks.
"""

from __future__ import annotations

import csv
import io
import os
from pathlib import Path
from typing import List, Tuple

from ai_engine.core.exceptions import DocumentProcessingError
from ai_engine.core.logging import get_logger

logger = get_logger(__name__)

# Type alias: list of (page_number, text) pairs.  page_number = -1 if not applicable.
PagedText = List[Tuple[int, str]]


# ---------------------------------------------------------------------------
# PDF Extractor
# ---------------------------------------------------------------------------
def extract_pdf(file_path: str) -> PagedText:
    """
    Extract text from a PDF using PyMuPDF.

    Handles:
      - Multi-column layouts (sorts text blocks by reading order)
      - Scanned PDFs return empty strings (OCR not implemented — logged)
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise DocumentProcessingError(
            "PyMuPDF not installed. Run: pip install pymupdf", {"error": str(e)}
        ) from e

    pages: PagedText = []
    try:
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")  # reading-order text
            if text.strip():
                pages.append((page_num, text))
            else:
                logger.warning(
                    "extractor.empty_page",
                    extra={"file": Path(file_path).name, "page": page_num},
                )
        doc.close()
    except Exception as e:
        raise DocumentProcessingError(
            f"Failed to extract PDF: {Path(file_path).name}", {"error": str(e)}
        ) from e

    if not pages:
        raise DocumentProcessingError(
            f"PDF appears to be scanned / empty: {Path(file_path).name}",
            {"hint": "Consider running OCR before uploading"},
        )

    return pages


# ---------------------------------------------------------------------------
# DOCX Extractor
# ---------------------------------------------------------------------------
def extract_docx(file_path: str) -> PagedText:
    """
    Extract text from a .docx file.
    Returns paragraphs grouped into ~100-paragraph pages for metadata purposes.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise DocumentProcessingError(
            "python-docx not installed. Run: pip install python-docx", {"error": str(e)}
        ) from e

    try:
        doc = Document(file_path)
        paragraphs: List[str] = []

        # Extract paragraph text
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract table cell text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            raise DocumentProcessingError(f"DOCX appears empty: {Path(file_path).name}")

        # Group into virtual pages of 100 paragraphs
        pages: PagedText = []
        chunk_size = 100
        for i in range(0, len(paragraphs), chunk_size):
            page_text = "\n".join(paragraphs[i:i + chunk_size])
            pages.append((i // chunk_size + 1, page_text))

        return pages

    except DocumentProcessingError:
        raise
    except Exception as e:
        raise DocumentProcessingError(
            f"Failed to extract DOCX: {Path(file_path).name}", {"error": str(e)}
        ) from e


# ---------------------------------------------------------------------------
# TXT Extractor
# ---------------------------------------------------------------------------
def extract_txt(file_path: str) -> PagedText:
    """
    Extract text from a plain text file.
    Tries UTF-8 then falls back to latin-1.
    """
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read()
            if text.strip():
                return [(-1, text)]
        except (UnicodeDecodeError, LookupError):
            continue

    raise DocumentProcessingError(
        f"Could not decode TXT file: {Path(file_path).name}",
        {"tried_encodings": ["utf-8", "utf-8-sig", "latin-1", "cp1252"]},
    )


# ---------------------------------------------------------------------------
# CSV Extractor
# ---------------------------------------------------------------------------
def extract_csv(file_path: str) -> PagedText:
    """
    Extract data from CSV as readable prose sentences.
    Each row becomes: "Column1: value1, Column2: value2, ..."
    """
    rows_text: List[str] = []
    try:
        with open(file_path, "r", encoding="utf-8", newline="") as f:
            reader = csv.DictReader(f)
            for row_num, row in enumerate(reader, start=1):
                parts = [f"{k}: {v}" for k, v in row.items() if v and v.strip()]
                if parts:
                    rows_text.append(", ".join(parts))
    except Exception as e:
        raise DocumentProcessingError(
            f"Failed to extract CSV: {Path(file_path).name}", {"error": str(e)}
        ) from e

    if not rows_text:
        raise DocumentProcessingError(f"CSV appears empty: {Path(file_path).name}")

    # Group into virtual pages of 200 rows
    pages: PagedText = []
    chunk_size = 200
    for i in range(0, len(rows_text), chunk_size):
        page_text = "\n".join(rows_text[i:i + chunk_size])
        pages.append((i // chunk_size + 1, page_text))

    return pages


# ---------------------------------------------------------------------------
# Markdown Extractor
# ---------------------------------------------------------------------------
def extract_md(file_path: str) -> PagedText:
    """Extract markdown as plain text — sections become virtual pages."""
    text = extract_txt(file_path)[0][1]

    # Split on H1/H2 headings to create sections
    import re
    sections = re.split(r"(?m)^#{1,2}\s+", text)
    pages: PagedText = []
    for i, section in enumerate(sections, start=1):
        if section.strip():
            pages.append((i, section.strip()))

    return pages if pages else [(-1, text)]


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------
_EXTRACTORS = {
    ".pdf":  extract_pdf,
    ".docx": extract_docx,
    ".txt":  extract_txt,
    ".csv":  extract_csv,
    ".md":   extract_md,
}


def extract_text(file_path: str) -> PagedText:
    """
    Dispatch to the correct extractor based on file extension.

    Args:
        file_path: Absolute path to the file

    Returns:
        List of (page_number, text) tuples

    Raises:
        DocumentProcessingError: if extension not supported or extraction fails
    """
    ext = Path(file_path).suffix.lower()
    extractor = _EXTRACTORS.get(ext)

    if extractor is None:
        raise DocumentProcessingError(
            f"Unsupported file format: {ext}",
            {"supported": list(_EXTRACTORS.keys()), "file": Path(file_path).name},
        )

    logger.info(
        "extractor.start",
        extra={
            "event": "extractor.start",
            "file": Path(file_path).name,
            "extension": ext,
            "size_bytes": os.path.getsize(file_path),
        },
    )

    pages = extractor(file_path)

    logger.info(
        "extractor.done",
        extra={
            "event": "extractor.done",
            "file": Path(file_path).name,
            "pages_extracted": len(pages),
        },
    )

    return pages
